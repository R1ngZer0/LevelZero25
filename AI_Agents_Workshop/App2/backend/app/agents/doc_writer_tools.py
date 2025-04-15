"""
Defines Pydantic AI tools used by the Document Writer Agent.
"""

from pydantic import BaseModel, Field
from typing import List
from pydantic_ai import Agent, tool
from app.llm_clients import get_llm_client
import os
from docx import Document
from docx.shared import Inches
from app.config import get_settings

# --- Tool 1: Outline Creator --- 

class OutlineTopic(BaseModel):
    """A single topic or section heading in the document outline."""
    heading: str = Field(..., description="The title or heading for this section.")
    # Optional: Add description/scope if needed
    # description: str | None = Field(default=None, description="Brief description of what this section should cover.")

class DocumentOutline(BaseModel):
    """Represents the structured outline for the document."""
    title: str = Field(..., description="The main title of the document.")
    topics: List[OutlineTopic] = Field(..., description="An ordered list of topics/sections for the document.")

@tool
def create_document_outline(prompt: str, context: str | None = None, mode: str = "cloud") -> DocumentOutline:
    """Generates a structured document outline based on a user prompt and optional context.

    Args:
        prompt: The user's request for the document (e.g., 'Write a report on AI ethics').
        context: Optional relevant context retrieved from the knowledge base.
        mode: The operational mode ('cloud' or 'local') for LLM selection.
    Returns:
        A DocumentOutline object containing the title and list of topics.
    """
    llm_client = get_llm_client(mode)
    
    system_prompt = (
        "You are an expert document planner. Based on the user's request and any provided context, "
        "create a logical and comprehensive outline for the document. Include a main title and a list of section headings. "
        "Focus on creating a clear structure. Output ONLY the DocumentOutline object."
    )
    
    combined_input = f"User Prompt: {prompt}"
    if context:
        combined_input += f"\n\nRelevant Context:\n{context}"

    outline_generator = Agent(
        llm=llm_client,
        system_prompt=system_prompt,
        output_model=DocumentOutline 
    )
    
    try:
        outline: DocumentOutline = outline_generator.invoke(combined_input)
        print(f"Generated outline for prompt: '{prompt[:50]}...'")
        return outline
    except Exception as e:
        print(f"Error generating document outline: {e}")
        # Fallback or raise error?
        # Re-raising allows the caller (agent) to handle it.
        # Consider adding FastAPI specific imports if HTTPException is intended here.
        # For now, re-raise the original exception for agent-level handling.
        raise RuntimeError(f"Failed to generate outline: {e}") from e

# --- Tool 2: Section Writer --- 

class SectionContent(BaseModel):
    """Represents the written content for a single document section."""
    heading: str = Field(..., description="The heading of the section this content belongs to.")
    content: str = Field(..., description="The written text content for the section.")

@tool
def write_document_section(
    topic_heading: str,
    full_prompt: str, 
    document_title: str, 
    previous_sections: List[SectionContent] = [],
    context: str | None = None, 
    mode: str = "cloud"
) -> SectionContent:
    """Writes the content for a specific section of a document based on its heading, 
    the overall prompt, optional context, and previously written sections.

    Args:
        topic_heading: The heading of the section to write.
        full_prompt: The original user prompt for the entire document.
        document_title: The main title of the document.
        previous_sections: Content of sections written before this one, for context.
        context: Optional relevant context retrieved from the knowledge base.
        mode: The operational mode ('cloud' or 'local') for LLM selection.
    Returns:
        A SectionContent object containing the heading and written content.
    """
    llm_client = get_llm_client(mode)
    
    system_prompt = (
        "You are an expert technical writer. Your task is to write a specific section of a document. "
        f"The document title is '{document_title}'. You need to write the section with the heading: '{topic_heading}'."
        "Use the overall user prompt and any provided context or previous sections to ensure coherence and accuracy. "
        "Focus on writing ONLY the content for the specified section. Do not include the heading itself in the output content. "
        "Output ONLY the SectionContent object."
        # NOTE: PydanticAI expects the *output* to be the model, so we can't just ask for raw text here.
        # We'll wrap the LLM call output in the SectionContent model.
    )
    
    combined_input = f"Overall Document Prompt: {full_prompt}\n"
    if context:
        combined_input += f"\nRelevant Knowledge Base Context:\n{context}\n"
    if previous_sections:
        combined_input += f"\n## Previously Written Sections (for context):\n"
        for sec in previous_sections:
            combined_input += f"### {sec.heading}\n{sec.content}\n\n"
    combined_input += f"\nWrite the content for the section: '{topic_heading}'"
    
    # Use Agent instead of PydanticAI
    section_writer_llm = Agent(
        llm=llm_client,
        system_prompt=system_prompt,
        output_model=SectionContent
    )

    try:
        # PydanticAI invoke expects the LLM to return JSON conforming to the model.
        # The prompt instructs the LLM to focus on the content, but PydanticAI handles the wrapping.
        # We might need to adjust the prompt if the LLM struggles to include the heading correctly.
        generated_section: SectionContent = section_writer_llm.invoke(combined_input)
        
        # Ensure the heading matches the requested topic (LLM might hallucinate)
        if generated_section.heading != topic_heading:
             print(f"Warning: LLM generated section with heading '{generated_section.heading}', expected '{topic_heading}'. Overwriting heading.")
             generated_section.heading = topic_heading
             
        print(f"Generated content for section: '{topic_heading}'")
        return generated_section
    except Exception as e:
        print(f"Error generating document section '{topic_heading}': {e}")
        raise RuntimeError(f"Failed to generate section '{topic_heading}': {e}") from e

# --- Tool 3: Report Writer --- 

settings = get_settings()

class ReportOutput(BaseModel):
    """Output schema for the report writer tool."""
    file_path: str = Field(..., description="The relative path within the knowledgebase where the final .docx file was saved.")
    full_text: str = Field(..., description="The full text content of the document, assembled from all sections.")

@tool
def format_and_save_report(
    document_title: str,
    sections: List[SectionContent],
    # Optional: Add parameters for file naming, etc.
    # base_filename: str | None = None 
) -> ReportOutput:
    """Formats the completed sections into an HTML string, then converts 
    and saves it as a .docx file in the knowledge base.

    Args:
        document_title: The main title of the document.
        sections: A list of SectionContent objects representing the approved sections.
        
    Returns:
        A ReportOutput object containing the path to the saved .docx file and the full text.
    """
    # 1. Assemble basic HTML and full text
    html_content = f"<h1>{document_title}</h1>\n"
    full_text_content = f"{document_title}\n\n"
    for section in sections:
        html_content += f"<h2>{section.heading}</h2>\n<p>{section.content}</p>\n"
        full_text_content += f"{section.heading}\n{section.content}\n\n"
        
    # 2. Convert HTML-like structure to DOCX using python-docx
    # (Note: python-docx doesn't directly parse HTML. We build the doc structure.)
    doc = Document()
    doc.add_heading(document_title, level=1)
    
    for section in sections:
        doc.add_heading(section.heading, level=2)
        # Add content as paragraphs
        # Basic handling: split by newline, treat each as paragraph
        paragraphs = section.content.strip().split('\n')
        for para in paragraphs:
            if para.strip(): # Avoid adding empty paragraphs
                doc.add_paragraph(para)
        doc.add_paragraph() # Add a space between sections

    # 3. Save the DOCX file
    knowledge_base_path = settings.KNOWLEDGE_BASE_PATH
    if not os.path.exists(knowledge_base_path):
        os.makedirs(knowledge_base_path)
        
    # Generate a safe filename (replace spaces, etc.)
    safe_title = "".join(c if c.isalnum() else '_' for c in document_title)
    output_filename = f"{safe_title[:50]}.docx" # Limit filename length
    absolute_file_path = os.path.join(knowledge_base_path, output_filename)
    
    # Ensure the filename is unique if needed (e.g., append timestamp or counter)
    # For simplicity, we overwrite for now.
    
    try:
        doc.save(absolute_file_path)
        relative_file_path = os.path.relpath(absolute_file_path, start=os.getcwd()) # Get path relative to workspace root
        print(f"Report saved successfully to: {relative_file_path}")
        return ReportOutput(file_path=relative_file_path, full_text=full_text_content.strip())
    except Exception as e:
        print(f"Error saving document '{output_filename}': {e}")
        raise RuntimeError(f"Failed to save report '{output_filename}': {e}") from e